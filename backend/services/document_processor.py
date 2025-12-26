"""
Unified Document Processing Service for Learning Mode
Handles extraction and processing of documents for both users and admins
Supports .txt, .docx, and .pdf files with vector embedding
"""

import os
import io
from typing import Optional, BinaryIO, Tuple
import logging

logger = logging.getLogger("uvicorn.error")

# Configuration
MAX_USER_UPLOAD_SIZE_MB = 200  # Total upload limit per user
MAX_SINGLE_FILE_SIZE_MB = 20   # Single file size limit
MAX_TEXT_SIZE_MB = 10  # Max extracted text size (prevent huge PDFs from crashing)
MAX_CHUNKS_PER_DOCUMENT = 5000  # Max chunks per document (prevent OOM)


def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from .txt file"""
    try:
        # Try UTF-8 first, fallback to other encodings
        try:
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return file_content.decode('latin-1')
            except:
                return file_content.decode('utf-8', errors='ignore')
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {e}")
        raise ValueError(f"Failed to extract text from TXT file: {e}")


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from .docx file"""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_content))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return '\n\n'.join(text_parts)
    except ImportError:
        raise ValueError("python-docx library not installed. Install with: pip install python-docx")
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise ValueError(f"Failed to extract text from DOCX file: {e}")


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from .pdf file"""
    try:
        import PyPDF2
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text_parts = []
        
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text.strip():
                text_parts.append(text)
        
        return '\n\n'.join(text_parts)
    except ImportError:
        raise ValueError("PyPDF2 library not installed. Install with: pip install PyPDF2")
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise ValueError(f"Failed to extract text from PDF file: {e}")


async def generate_document_summary(content: str, file_name: str) -> Tuple[str, dict]:
    """
    Generate a summary of document content using LLM.
    
    Args:
        content: Full text content of the document
        file_name: Name of the file for metadata
    
    Returns:
        Tuple of (summary_text, metadata_dict)
    """
    from services.llm_provider import get_llm_provider
    
    # Truncate content to prevent token overflow (approx 8000 tokens = 32000 chars)
    max_content_chars = 32000
    truncated = content[:max_content_chars] if len(content) > max_content_chars else content
    was_truncated = len(content) > max_content_chars
    
    provider = get_llm_provider()
    
    prompt = f"""Hãy tóm tắt nội dung tài liệu sau đây một cách ngắn gọn, súc tích (tối đa 500 từ).
Tập trung vào các ý chính, khái niệm quan trọng, và thông tin hữu ích cho việc học tập.

Tên tài liệu: {file_name}

Nội dung:
{truncated}

{"[Lưu ý: Tài liệu đã được cắt ngắn do quá dài]" if was_truncated else ""}

Tóm tắt:"""

    try:
        messages = [
            {"role": "system", "content": "Bạn là trợ lý tóm tắt tài liệu học tập. Hãy viết tóm tắt ngắn gọn, dễ hiểu."},
            {"role": "user", "content": prompt}
        ]
        
        response = await provider.chat(messages=messages, temperature=0.3)
        
        summary = ""
        if response and isinstance(response, dict):
            candidates = response.get("candidates", [])
            if candidates and isinstance(candidates[0], dict):
                content_obj = candidates[0].get("content", {})
                parts = content_obj.get("parts", [])
                if parts and isinstance(parts[0], dict):
                    summary = parts[0].get("text", "")
        
        if not summary:
            # Fallback: Use first 500 chars as summary
            summary = truncated[:500] + "..." if len(truncated) > 500 else truncated
        
        metadata = {
            "file_name": file_name,
            "content_length": len(content),
            "was_truncated": was_truncated,
            "summary_generated": bool(summary),
            "method": "llm" if summary else "fallback"
        }
        
        return summary, metadata
        
    except Exception as e:
        logger.error(f"Error generating document summary: {e}")
        # Fallback to simple truncation
        fallback_summary = truncated[:500] + "..." if len(truncated) > 500 else truncated
        return fallback_summary, {
            "file_name": file_name,
            "content_length": len(content),
            "error": str(e),
            "method": "fallback"
        }


def extract_document_text(file_content: bytes, file_type: str) -> str:
    """
    Extract text from document based on file type
    
    Args:
        file_content: Binary content of the file
        file_type: File extension ('txt', 'docx', 'pdf')
    
    Returns:
        Extracted text content (with NUL characters removed)
    
    Raises:
        ValueError: If file type is not supported or extraction fails
    """
    file_type = file_type.lower().strip('.')
    
    if file_type == 'txt':
        text = extract_text_from_txt(file_content)
    elif file_type == 'docx':
        text = extract_text_from_docx(file_content)
    elif file_type == 'pdf':
        text = extract_text_from_pdf(file_content)
    else:
        raise ValueError(f"Unsupported file type: {file_type}. Supported types: txt, docx, pdf")
    
    # Remove NUL characters that cause PostgreSQL errors
    text = text.replace('\x00', '')
    
    # Check extracted text size to prevent memory issues
    text_size_mb = len(text.encode('utf-8')) / (1024 * 1024)
    if text_size_mb > MAX_TEXT_SIZE_MB:
        logger.warning(f"Extracted text too large: {text_size_mb:.2f}MB, truncating to {MAX_TEXT_SIZE_MB}MB")
        # Truncate to max size (keeping first portion)
        max_bytes = int(MAX_TEXT_SIZE_MB * 1024 * 1024)
        text_bytes = text.encode('utf-8')[:max_bytes]
        text = text_bytes.decode('utf-8', errors='ignore')
    
    return text


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
    """
    Split text into overlapping chunks for vector storage
    
    Args:
        text: Text to split
        chunk_size: Maximum characters per chunk
        overlap: Number of characters to overlap between chunks
    
    Returns:
        List of text chunks (with NUL characters removed)
    """
    # Remove NUL characters that cause PostgreSQL errors
    text = text.replace('\x00', '')
    
    if not text or not text.strip():
        return []
    
    chunks = []
    start = 0
    text_length = len(text)
    
    while start < text_length:
        # Safety check: prevent infinite chunks
        if len(chunks) >= MAX_CHUNKS_PER_DOCUMENT:
            logger.warning(f"Reached max chunks limit ({MAX_CHUNKS_PER_DOCUMENT}), stopping chunking")
            break
        
        end = start + chunk_size
        
        # If not at the end, try to break at a sentence or paragraph boundary
        if end < text_length:
            # Look for paragraph break first
            break_pos = text.rfind('\n\n', start, end)
            if break_pos == -1:
                # Look for sentence break
                break_pos = text.rfind('. ', start, end)
            if break_pos == -1:
                # Look for any space
                break_pos = text.rfind(' ', start, end)
            if break_pos > start:
                end = break_pos + 1
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        # Move start position with overlap
        start = end - overlap if end < text_length else text_length
    
    return chunks


def validate_document_file(
    filename: str, 
    file_size: int, 
    max_size_mb: int = MAX_SINGLE_FILE_SIZE_MB,
    is_admin: bool = False
) -> tuple[bool, Optional[str]]:
    """
    Validate document file before processing
    
    Args:
        filename: Original filename
        file_size: File size in bytes
        max_size_mb: Maximum allowed file size in MB
        is_admin: Whether this is an admin upload (no size limit)
    
    Returns:
        Tuple of (is_valid, error_message)
    """
    # Check file extension
    allowed_extensions = {'.txt', '.docx', '.pdf'}
    file_ext = os.path.splitext(filename)[1].lower()
    
    if file_ext not in allowed_extensions:
        return False, f"File type not allowed. Allowed types: {', '.join(allowed_extensions)}"
    
    # Check file size (skip for admin)
    if not is_admin:
        max_size_bytes = max_size_mb * 1024 * 1024
        if file_size > max_size_bytes:
            return False, f"File too large. Maximum size: {max_size_mb}MB"
    
    if file_size == 0:
        return False, "File is empty"
    
    return True, None


def check_user_storage_limit(db, user_id: int, new_file_size: int) -> tuple[bool, Optional[str]]:
    """
    Check if user has exceeded storage limit
    
    Args:
        db: Database session
        user_id: User ID
        new_file_size: Size of new file to upload (bytes)
    
    Returns:
        Tuple of (can_upload, error_message)
    """
    from db import models
    
    # Get total size of user's documents
    user_docs = db.query(models.Document).filter(
        models.Document.user_id == user_id
    ).all()
    
    total_size = sum(doc.file_size for doc in user_docs)
    max_size_bytes = MAX_USER_UPLOAD_SIZE_MB * 1024 * 1024
    
    if total_size + new_file_size > max_size_bytes:
        used_mb = total_size / (1024 * 1024)
        return False, f"Storage limit exceeded. You have used {used_mb:.1f}MB / {MAX_USER_UPLOAD_SIZE_MB}MB"
    
    return True, None


def check_user_storage_limit_batch(db, user_id: int, new_files_total_size: int) -> tuple[bool, Optional[str]]:
    """
    Check if user has exceeded storage limit for batch upload
    
    Args:
        db: Database session
        user_id: User ID
        new_files_total_size: Total size of all new files to upload (bytes)
    
    Returns:
        Tuple of (can_upload, error_message)
    """
    from db import models
    
    # Get total size of user's existing documents
    user_docs = db.query(models.Document).filter(
        models.Document.user_id == user_id
    ).all()
    
    existing_total_size = sum(doc.file_size for doc in user_docs)
    max_size_bytes = MAX_USER_UPLOAD_SIZE_MB * 1024 * 1024
    
    if existing_total_size + new_files_total_size > max_size_bytes:
        used_mb = existing_total_size / (1024 * 1024)
        new_mb = new_files_total_size / (1024 * 1024)
        return False, f"Storage limit exceeded. Current: {used_mb:.1f}MB, Upload: {new_mb:.1f}MB, Limit: {MAX_USER_UPLOAD_SIZE_MB}MB"
    
    return True, None


async def summarize_with_llm(
    full_text: str,
    file_name: str,
    context: str = "educational document"
) -> Tuple[str, dict]:
    """
    Use LLM to extract key information from document
    
    Args:
        full_text: Full text content
        file_name: Original filename
        context: Context description
    
    Returns:
        Tuple of (summary, metadata)
    """
    from services.llm_provider import get_llm_provider
    
    estimated_tokens = len(full_text) // 4
    target_summary_tokens = 2000
    
    prompt = f"""Bạn là chuyên gia phân tích tài liệu giáo dục. Hãy trích xuất và tóm tắt thông tin quan trọng từ tài liệu sau.

**Tài liệu**: {file_name}
**Bối cảnh**: {context}
**Độ dài**: ~{estimated_tokens} tokens

**YÊU CẦU**:
1. Trích xuất các khái niệm, công thức, định lý quan trọng
2. Tóm tắt các phần kiến thức chính
3. Ghi chú các ví dụ và bài tập tiêu biểu
4. Sử dụng Markdown với headers rõ ràng
5. Độ dài tối đa: ~{target_summary_tokens * 4} ký tự

**NỘI DUNG**:
{full_text[:50000]}
{'... (còn ' + str(len(full_text) - 50000) + ' ký tự)' if len(full_text) > 50000 else ''}

Hãy trích xuất thông tin quan trọng:"""

    try:
        llm = get_llm_provider()
        messages = [
            {"role": "system", "content": "Bạn là chuyên gia phân tích tài liệu giáo dục."},
            {"role": "user", "content": prompt}
        ]
        response = await llm.chat(messages, temperature=0.3)
        
        # Extract text from response
        summary = ""
        if isinstance(response, dict):
            if "candidates" in response:
                candidates = response.get("candidates", [])
                if candidates and isinstance(candidates[0], dict):
                    content = candidates[0].get("content", {})
                    if isinstance(content, dict):
                        parts = content.get("parts", [])
                        if parts and isinstance(parts[0], dict):
                            summary = parts[0].get("text", "")
            elif "choices" in response:
                choices = response.get("choices", [])
                if choices:
                    msg = choices[0].get("message", {})
                    summary = msg.get("content", "") if isinstance(msg, dict) else ""
        elif isinstance(response, str):
            summary = response
        
        if not summary:
            # Fallback to truncation
            summary = full_text[:target_summary_tokens * 4]
        
        metadata = {
            "original_length": len(full_text),
            "summary_length": len(summary),
            "compression_ratio": round(len(full_text) / len(summary), 2) if summary else 0,
            "method": "llm_summary" if summary else "truncation"
        }
        
        return summary, metadata
    
    except Exception as e:
        logger.error(f"LLM summarization failed: {e}")
        # Fallback
        summary = full_text[:target_summary_tokens * 4]
        metadata = {
            "original_length": len(full_text),
            "summary_length": len(summary),
            "method": "truncation_fallback",
            "error": str(e)
        }
        return summary, metadata


async def process_document(file_content: bytes, filename: str) -> str:
    """
    Main function to process uploaded documents
    Returns extracted text content
    """
    try:
        file_ext = '.' + filename.split('.')[-1].lower() if '.' in filename else ''
        
        if file_ext == '.txt':
            text_content = extract_text_from_txt(file_content)
        elif file_ext == '.docx':
            text_content = extract_text_from_docx(file_content)
        elif file_ext == '.pdf':
            text_content = extract_text_from_pdf(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        # Basic validation
        if not text_content.strip():
            raise ValueError("No text content extracted from file")
        
        logger.info(f"Processed document: {filename} ({len(text_content)} chars)")
        return text_content
        
    except Exception as e:
        logger.error(f"Document processing failed for {filename}: {e}")
        raise ValueError(f"Failed to process document: {e}")


async def process_uploaded_document(
    file_bytes: bytes,
    file_name: str,
    file_type: str,
    structure_name: str = ""
) -> tuple:
    """
    Process an uploaded document for custom teaching structures.
    Returns: (original_content, extracted_summary, metadata)
    
    This replaces the old document_extractor.process_uploaded_document function.
    """
    try:
        # Extract text based on file type
        file_ext = f".{file_type.lower()}"
        
        if file_ext == '.txt':
            original_content = extract_text_from_txt(file_bytes)
        elif file_ext in ['.docx', '.doc']:
            original_content = extract_text_from_docx(file_bytes)
        elif file_ext == '.pdf':
            original_content = extract_text_from_pdf(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        if not original_content.strip():
            raise ValueError("No text content extracted from file")
        
        # Generate summary using LLM
        summary, metadata = await generate_document_summary(original_content, file_name)
        
        # Add structure context to metadata
        metadata['structure_name'] = structure_name
        metadata['original_length'] = len(original_content)
        metadata['summary_length'] = len(summary)
        metadata['compression_ratio'] = round(len(summary) / max(len(original_content), 1), 2)
        
        logger.info(f"Processed uploaded document: {file_name} ({len(original_content)} -> {len(summary)} chars)")
        
        return original_content, summary, metadata
        
    except Exception as e:
        logger.error(f"Failed to process uploaded document {file_name}: {e}")
        raise ValueError(f"Document processing failed: {e}")
