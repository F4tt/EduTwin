"""
Document Extraction Service
Extracts and summarizes content from PDF, DOCX, and TXT files for AI context.
Uses LLM to intelligently extract key information while reducing token usage.
"""

import os
import io
import logging
from typing import Optional, Tuple
import PyPDF2
from docx import Document as DocxDocument

from services.llm_provider import get_llm_provider

logger = logging.getLogger(__name__)

# Maximum characters to extract from original document
MAX_ORIGINAL_CONTENT_CHARS = 100000  # ~25k tokens

# Target summary length (in tokens, approximately)
TARGET_SUMMARY_TOKENS = 2000  # Optimized for context inclusion


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file."""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        
        for page in pdf_reader.pages:
            try:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            except Exception as e:
                logger.warning(f"Failed to extract text from PDF page: {e}")
                continue
        
        full_text = "\n\n".join(text_parts)
        return full_text[:MAX_ORIGINAL_CONTENT_CHARS]
    
    except Exception as e:
        logger.error(f"Failed to extract PDF: {e}")
        raise ValueError(f"Không thể đọc file PDF: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file."""
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        full_text = "\n\n".join(text_parts)
        return full_text[:MAX_ORIGINAL_CONTENT_CHARS]
    
    except Exception as e:
        logger.error(f"Failed to extract DOCX: {e}")
        raise ValueError(f"Không thể đọc file DOCX: {str(e)}")


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from TXT file."""
    try:
        # Try UTF-8 first, fallback to other encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                text = file_bytes.decode(encoding)
                return text[:MAX_ORIGINAL_CONTENT_CHARS]
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Không thể xác định encoding của file text")
    
    except Exception as e:
        logger.error(f"Failed to extract TXT: {e}")
        raise ValueError(f"Không thể đọc file TXT: {str(e)}")


def extract_document_content(file_bytes: bytes, file_type: str) -> str:
    """Extract text content from document based on file type."""
    file_type = file_type.lower().replace('.', '')
    
    extractors = {
        'pdf': extract_text_from_pdf,
        'docx': extract_text_from_docx,
        'doc': extract_text_from_docx,  # Try DOCX extractor for DOC
        'txt': extract_text_from_txt,
    }
    
    extractor = extractors.get(file_type)
    if not extractor:
        raise ValueError(f"Không hỗ trợ định dạng file: {file_type}. Chỉ hỗ trợ PDF, DOCX, TXT")
    
    return extractor(file_bytes)


async def summarize_educational_document(
    full_text: str,
    structure_name: str,
    file_name: str
) -> Tuple[str, dict]:
    """
    Use LLM to extract key educational information from document.
    Returns: (extracted_summary, metadata)
    
    This creates a compact, structured summary optimized for AI context.
    """
    if not full_text or not full_text.strip():
        raise ValueError("Document is empty")
    
    # Calculate approximate token count (rough estimate: 4 chars = 1 token)
    estimated_tokens = len(full_text) // 4
    
    prompt = f"""Bạn là chuyên gia phân tích tài liệu giáo dục. Hãy trích xuất và tóm tắt thông tin quan trọng từ tài liệu sau để làm ngữ cảnh cho hệ thống AI phân tích học tập.

**Tài liệu**: {file_name}
**Cấu trúc giảng dạy**: {structure_name}
**Độ dài tài liệu**: ~{estimated_tokens} tokens

**YÊU CẦU TRÍCH XUẤT**:
1. **Thang điểm và phân loại**: Các mức điểm, tiêu chí đánh giá học lực
2. **Các mốc quan trọng**: Điểm chuẩn, mục tiêu theo từng cấp độ
3. **Xu hướng và phân tích**: Cách đánh giá tiến bộ, mất cân bằng
4. **Các tổ hợp môn/khối thi** (nếu có): Danh sách và đặc điểm
5. **Lời khuyên theo bối cảnh**: Phương pháp học theo mức điểm
6. **Dấu hiệu cần lưu ý**: Các chỉ báo tích cực/tiêu cực

**ĐỊNH DẠNG ĐẦU RA**:
- Sử dụng Markdown với headers rõ ràng
- Bullet points ngắn gọn, súc tích
- Ưu tiên thông tin số liệu cụ thể (điểm số, phần trăm, v.v.)
- Loại bỏ thông tin dư thừa, lặp lại
- Độ dài tối đa: ~{TARGET_SUMMARY_TOKENS * 4} ký tự (~{TARGET_SUMMARY_TOKENS} tokens)

**TÀI LIỆU GỐC**:
{full_text[:50000]}  
{'... (tài liệu còn ' + str(len(full_text) - 50000) + ' ký tự nữa)' if len(full_text) > 50000 else ''}

Hãy trích xuất thông tin quan trọng theo định dạng yêu cầu:"""

    try:
        llm = get_llm_provider()
        
        # Use chat method which is available in LLMProvider
        messages = [
            {"role": "system", "content": "Bạn là chuyên gia phân tích tài liệu giáo dục. Hãy trích xuất thông tin quan trọng một cách súc tích."},
            {"role": "user", "content": prompt}
        ]
        response = await llm.chat(messages, temperature=0.3)
        
        # Extract text from response (handle different response formats)
        summary = ""
        if isinstance(response, dict):
            # Try various response formats
            if "candidates" in response:
                candidates = response.get("candidates", [])
                if candidates and isinstance(candidates[0], dict):
                    content = candidates[0].get("content", {})
                    if isinstance(content, dict):
                        parts = content.get("parts", [])
                        if parts and isinstance(parts[0], dict):
                            summary = parts[0].get("text", "")
            elif "choices" in response:
                choices = response.get("choices", [])
                if choices:
                    msg = choices[0].get("message", {})
                    summary = msg.get("content", "") if isinstance(msg, dict) else ""
            
            if not summary:
                # Fallback: scan for text in response
                def _scan(obj):
                    if isinstance(obj, str) and len(obj) > 100:
                        return obj
                    if isinstance(obj, dict):
                        for v in obj.values():
                            res = _scan(v)
                            if res:
                                return res
                    if isinstance(obj, list):
                        for v in obj:
                            res = _scan(v)
                            if res:
                                return res
                    return None
                summary = _scan(response) or ""
        elif isinstance(response, str):
            summary = response
        
        if not summary:
            raise ValueError("Empty response from LLM")
        
        metadata = {
            "original_length_chars": len(full_text),
            "original_estimated_tokens": estimated_tokens,
            "summary_length_chars": len(summary),
            "summary_estimated_tokens": len(summary) // 4,
            "compression_ratio": round(len(full_text) / len(summary), 2) if summary else 0,
            "extraction_method": "llm_summary",
            "structure_name": structure_name,
            "file_name": file_name
        }
        
        logger.info(f"Document summarized: {file_name}, compression ratio: {metadata['compression_ratio']}x")
        
        return summary, metadata
    
    except Exception as e:
        logger.error(f"Failed to summarize document: {e}")
        # Fallback: truncate to target length
        fallback_summary = full_text[:TARGET_SUMMARY_TOKENS * 4]
        metadata = {
            "original_length_chars": len(full_text),
            "summary_length_chars": len(fallback_summary),
            "extraction_method": "truncation_fallback",
            "error": str(e)
        }
        return fallback_summary, metadata


async def process_uploaded_document(
    file_bytes: bytes,
    file_name: str,
    file_type: str,
    structure_name: str
) -> Tuple[str, str, dict]:
    """
    Complete document processing pipeline.
    
    Returns:
        - original_content: Full extracted text (for reference)
        - extracted_summary: LLM-optimized summary
        - metadata: Processing metadata
    """
    # Step 1: Extract full text
    original_content = extract_document_content(file_bytes, file_type)
    
    if not original_content or not original_content.strip():
        raise ValueError("File không có nội dung hoặc không thể đọc được")
    
    # Step 2: Summarize with LLM
    extracted_summary, metadata = await summarize_educational_document(
        original_content,
        structure_name,
        file_name
    )
    
    return original_content, extracted_summary, metadata


def get_structure_knowledge_context(db, structure_id: int) -> str:
    """
    Load and combine extracted summaries from all documents of a structure
    into a single knowledge context string for AI.
    
    Args:
        db: Database session
        structure_id: ID of the teaching structure
    
    Returns:
        Combined knowledge base string
    """
    from db import models
    
    documents = (
        db.query(models.CustomStructureDocument)
        .filter(models.CustomStructureDocument.structure_id == structure_id)
        .order_by(models.CustomStructureDocument.uploaded_at.desc())
        .all()
    )
    
    if not documents:
        logger.info(f"No documents found for structure ID {structure_id}")
        return ""
    
    context_parts = ["# KIẾN THỨC THAM KHẢO CHO CẤU TRÚC GIẢNG DẠY\n"]
    
    for i, doc in enumerate(documents, 1):
        context_parts.append(f"## Tài liệu {i}: {doc.file_name}\n")
        context_parts.append(doc.extracted_summary)
        context_parts.append("\n---\n")
    
    total_chars = sum(len(doc.extracted_summary) for doc in documents)
    estimated_tokens = total_chars // 4
    
    logger.info(f"Built knowledge context from {len(documents)} documents, ~{estimated_tokens} tokens total")
    
    return "\n".join(context_parts)
